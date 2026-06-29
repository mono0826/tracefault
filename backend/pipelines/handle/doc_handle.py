import os

from .base_handle import BaseFileHandler
from ..models import Document, FileType


class DocFileHandler(BaseFileHandler):
    def process(self, file: str) -> Document:
        """
        读取旧版Word文档(.doc)
        首先尝试使用Windows特有的方法，如果失败则使用跨平台的方法
        """
        content = ""

        # 方法1: 尝试使用win32com (仅Windows)
        try:
            import win32com.client

            print(f"尝试使用win32com读取.doc文件: {os.path.basename(file)}")
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False

            doc_abs_path = os.path.abspath(file)
            doc = word.Documents.Open(doc_abs_path)
            content = doc.Content.Text
            doc.Close()
            word.Quit()

            if content and content.strip():
                print(f"使用win32com成功读取.doc文件")
                return Document(
                    file_name=os.path.basename(file),
                    file_type=FileType.DOC,
                    file_path=file,
                    file_size=os.path.getsize(file),
                    file_content=content,
                    status="success"
                )
        except ImportError:
            print("win32com不可用，这不是Windows系统")
        except Exception as e:
            print(f"使用win32com读取.doc失败: {str(e)}")

        # 方法2: 尝试使用textract (跨平台)
        try:
            import textract
            print(f"尝试使用textract读取.doc文件: {os.path.basename(file)}")
            content = textract.process(file).decode('utf-8')

            if content and content.strip():
                print(f"使用textract成功读取.doc文件")
                return Document(
                    file_name=os.path.basename(file),
                    file_type=FileType.DOC,
                    file_path=file,
                    file_size=os.path.getsize(file),
                    file_content=content,
                    status="success"
                )
        except ImportError:
            print("textract不可用，请安装: pip install textract")
        except Exception as e:
            print(f"使用textract读取.doc失败: {str(e)}")

        # 方法3: 尝试使用python-docx (不完全兼容.doc，但有时可以部分读取)
        try:
            from docx import Document as DocxDocument
            print(f"尝试使用python-docx读取.doc文件: {os.path.basename(file)}")
            doc = DocxDocument(file)
            full_text = []
            for para in doc.paragraphs:
                full_text.append(para.text)
            content = '\n'.join(full_text)

            if content and content.strip():
                print(f"使用python-docx部分读取.doc文件成功")
                return Document(
                    file_name=os.path.basename(file),
                    file_type=FileType.DOC,
                    file_path=file,
                    file_size=os.path.getsize(file),
                    file_content=content,
                    status="success"
                )
        except ImportError:
            print("python-docx不可用，请安装: pip install python-docx")
        except Exception as e:
            print(f"尝试使用python-docx读取.doc失败: {str(e)}")

        # 所有方法都失败，返回警告信息
        warning_msg = f"[警告: 无法读取.doc文件 {os.path.basename(file)}，请安装相关依赖或转换为.docx格式]"
        print(warning_msg)
        return Document(
            file_name=os.path.basename(file),
            file_type=FileType.DOC,
            file_path=file,
            file_size=os.path.getsize(file),
            file_content=warning_msg,
            status="failed"
        )

if __name__ == "__main__":
    import sys
    from pathlib import Path

    project_root = str(Path(__file__).resolve().parent.parent.parent)
    if project_root not in sys.path:
        sys.path.insert(0, project_root)

    data_dir = Path(__file__).resolve().parent.parent.parent.parent / "data" / "test"
    test_file = data_dir / "test.doc"

    handler = DocFileHandler()

    if test_file.exists():
        result = handler.process(str(test_file))
        print(f"\n=== DocFileHandler 测试 ===")
        print(f"文件名: {result.file_name}")
        print(f"文件类型: {result.file_type}")
        print(f"文件路径: {result.file_path}")
        print(f"文件大小: {result.file_size} 字节")
        print(f"状态: {result.status}")
        print(f"内容预览:\n{result.file_content[:500]}...")
    else:
        print(f"\n测试文件不存在: {test_file}，跳过 DocFileHandler 测试")
        print("提示: 可以将 .doc 文件放入 data/test/ 目录后测试")