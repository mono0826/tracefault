import os
from docx import Document as DocxDocument

from .base_handle import BaseFileHandler
from ..models import Document, FileType


class DocxFileHandler(BaseFileHandler):
    def process(self, file: str) -> Document:
        """读取Word文档(.docx)"""
        try:
            doc = DocxDocument(file)
            full_text = []
            for para in doc.paragraphs:
                full_text.append(para.text)
            content = '\n'.join(full_text)

            return Document(
                file_name=file,
                file_type=FileType.DOCX,
                file_path=os.path.basename(file),
                file_size=os.path.getsize(file),
                file_content=content,
                status="success"
            )
        except Exception as e:
            print(f"读取Word文档(.docx) {os.path.basename(file)} 失败: {str(e)}")
            return Document(
                file_name=file,
                file_type=FileType.DOCX,
                file_path=os.path.basename(file),
                file_size=os.path.getsize(file),
                file_content=f"[无法读取Word文档内容: {str(e)}]",
                status="failed"
            )


if __name__ == "__main__":
    import sys
    from pathlib import Path

    project_root = str(Path(__file__).resolve().parent.parent.parent)
    if project_root not in sys.path:
        sys.path.insert(0, project_root)

    data_dir = Path(__file__).resolve().parent.parent.parent.parent / "data" / "test"
    test_file = data_dir / "test.docx"

    handler = DocxFileHandler()

    if test_file.exists():
        result = handler.process(str(test_file))
        print(f"\n=== DocxFileHandler 测试 ===")
        print(f"文件名: {result.file_name}")
        print(f"文件类型: {result.file_type}")
        print(f"文件路径: {result.file_path}")
        print(f"文件大小: {result.file_size} 字节")
        print(f"状态: {result.status}")
        print(f"内容预览:\n{result.file_content[:500]}...")
    else:
        print(f"\n测试文件不存在: {test_file}")
